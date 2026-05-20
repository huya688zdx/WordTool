"""Generate test DOCX files for development and testing."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


def create_simple_docx(output_path: Path):
    """Create a simple DOCX with 3 paragraphs."""
    doc = Document()

    doc.add_heading("Test Document", level=1)

    doc.add_paragraph("This is the first paragraph with simple text.")
    doc.add_paragraph("This is the second paragraph that contains some important information.")
    doc.add_paragraph("This is the third paragraph with additional details.")

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_highlighted_docx(output_path: Path):
    """Create a DOCX with highlighted text."""
    doc = Document()

    doc.add_heading("Highlighted Document", level=1)

    # Paragraph with yellow highlight
    p1 = doc.add_paragraph()
    run1 = p1.add_run("This text has ")
    run2 = p1.add_run("yellow highlighted")
    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run3 = p1.add_run(" content.")

    # Paragraph with red text (simulating change)
    p2 = doc.add_paragraph()
    run4 = p2.add_run("This section has ")
    run5 = p2.add_run("red colored text")
    run5.font.color.rgb = RGBColor(255, 0, 0)
    run6 = p2.add_run(" for emphasis.")

    # Normal paragraph
    doc.add_paragraph("This is a normal paragraph without any highlighting.")

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_multiline_docx(output_path: Path):
    """Create a DOCX with long paragraphs that span multiple lines."""
    doc = Document()

    doc.add_heading("Multi-page Document", level=1)

    # Add many paragraphs to create a multi-page document
    for i in range(50):
        doc.add_paragraph(
            f"This is paragraph number {i + 1}. "
            f"It contains enough text to span multiple lines in the document. "
            f"The purpose of this paragraph is to test the text anchor mapping "
            f"algorithm with longer text segments that may cross page boundaries."
        )

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_table_docx(output_path: Path):
    """Create a DOCX with tables."""
    doc = Document()

    doc.add_heading("Table Document", level=1)

    doc.add_paragraph("This paragraph comes before the table.")

    # Create a table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    # Fill table cells
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell ({i+1}, {j+1})"

    doc.add_paragraph("This paragraph comes after the table.")

    doc.save(str(output_path))
    print(f"Created: {output_path}")


if __name__ == "__main__":
    fixtures_dir = Path(__file__).parent.parent / "tests" / "fixtures"
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    create_simple_docx(fixtures_dir / "simple.docx")
    create_highlighted_docx(fixtures_dir / "highlighted.docx")
    create_multiline_docx(fixtures_dir / "multiline.docx")
    create_table_docx(fixtures_dir / "table.docx")

    print("\nAll test fixtures created successfully!")
